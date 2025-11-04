"""
RAG Training Service
Comprehensive RAG training implementation with dataset refinement and testing
"""

import os
import json
import time
from typing import List, Dict, Any, Optional, Tuple
from datetime import datetime

# Import services
from ..dataset_processing.dataset_refiner import get_dataset_refiner
from ..metrics.training_metrics_collector import TrainingMetricsCollector
from app.services.chromadb_service import ChromaDBService
from app.services.external_api_service import ExternalAPIService
from services.training_websocket_service import get_training_websocket_service

# File processing imports
try:
    import docx
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    print("⚠️ python-docx not available - DOCX files will be skipped")

try:
    import pypdf
    PDF_AVAILABLE = True
    PDF_LIB = 'pypdf'
except ImportError:
    try:
        import PyPDF2
        PDF_AVAILABLE = True
        PDF_LIB = 'PyPDF2'
    except ImportError:
        PDF_AVAILABLE = False
        PDF_LIB = None
        print("⚠️ PyPDF2/pypdf not available - PDF files will be skipped")


class RAGTrainingService:
    """
    Comprehensive RAG training service with dataset refinement and testing
    """
    
    def __init__(self):
        self.dataset_refiner = get_dataset_refiner()
        self.metrics_collector = TrainingMetricsCollector()
        self.chromadb_service = ChromaDBService()
        self.external_api_service = ExternalAPIService()
        
    def train_minion_with_rag(self, 
                            job_id: int,
                            minion_id: int,
                            user_id: int,
                            datasets: List[Dict[str, Any]],
                            rag_config: Dict[str, Any],
                            minion_config: Dict[str, Any]) -> Dict[str, Any]:
        """
        Complete RAG training pipeline with before/after metrics
        
        Args:
            job_id: Training job ID
            minion_id: Minion ID to train
            user_id: User ID
            datasets: List of datasets to use
            rag_config: RAG configuration
            minion_config: Minion configuration
            
        Returns:
            Training results and metrics
        """
        print(f"\n{'='*60}")
        print(f"🚀 Starting RAG Training for Minion {minion_id}")
        print(f"{'='*60}\n")
        
        # Initialize metrics collection
        self.metrics_collector.start_collection(job_id, minion_id)
        
        # Import TrainingService for progress updates
        from app.services.training_service import TrainingService
        training_service = TrainingService()
        
        # Total steps: 9 (0-8)
        TOTAL_STEPS = 9
        STEP_PROGRESS = 100.0 / TOTAL_STEPS
        
        try:
            # Websockets disabled to prevent Flask-SocketIO conflicts
            # websocket_service = get_training_websocket_service()
            
            # Step 0: Capture BEFORE metrics
            print("📊 Step 0: Capturing BEFORE training metrics")
            before_metrics = self._capture_before_metrics(minion_id, minion_config)
            self.metrics_collector.record_before_metrics(before_metrics)
            training_service.update_training_job_progress(job_id, STEP_PROGRESS * 1, current_step=0)
            
            # Step 1: Dataset Loading and Refinement
            print("📊 Step 1: Dataset Loading and Refinement")
            refined_datasets, refinement_stats = self._refine_datasets(datasets, rag_config)
            self.metrics_collector.record_dataset_stats(
                len(datasets), 
                len(refined_datasets), 
                refinement_stats
            )
            training_service.update_training_job_progress(job_id, STEP_PROGRESS * 2, current_step=1)
            
            # Step 2: Knowledge Base Creation
            print("🗄️ Step 2: Knowledge Base Creation")
            collection_name = self._create_knowledge_base(job_id, refined_datasets, rag_config)
            self.metrics_collector.record_knowledge_base_stats(
                len(refined_datasets), 
                collection_name
            )
            training_service.update_training_job_progress(job_id, STEP_PROGRESS * 3, current_step=2)
            
            # Step 3: Embedding Creation
            print("🔮 Step 3: Embedding Creation")
            embedding_stats = self._create_embeddings(collection_name, refined_datasets, rag_config)
            training_service.update_training_job_progress(job_id, STEP_PROGRESS * 4, current_step=3)
            
            # Step 4: Minion Update
            print("🤖 Step 4: Minion Update")
            minion_update_result = self._update_minion_with_rag(minion_id, collection_name, rag_config)
            
            # ✅ CRITICAL: Check if minion update failed - STOP TRAINING IMMEDIATELY
            if not minion_update_result.get('rag_config_applied', False):
                error_msg = minion_update_result.get('error', 'Unknown minion update error')
                print(f"\n{'='*60}")
                print(f"❌ RAG TRAINING FAILED: {error_msg}")
                print(f"{'='*60}")
                raise Exception(f"Minion update failed: {error_msg}")
            
            print(f"   ✅ Minion update successful - continuing training...")
            training_service.update_training_job_progress(job_id, STEP_PROGRESS * 5, current_step=4)
            
            # Step 5: Training Validation
            print("✅ Step 5: Training Validation")
            validation_results = self._validate_training(minion_id, collection_name, rag_config)
            self.metrics_collector.record_validation_stats(validation_results)
            training_service.update_training_job_progress(job_id, STEP_PROGRESS * 6, current_step=5)
            
            # Step 6: Testing
            print("🧪 Step 6: Testing")
            test_results = self._test_minion_performance(minion_id, collection_name, rag_config)
            training_service.update_training_job_progress(job_id, STEP_PROGRESS * 7, current_step=6)
            
            # Step 7: Capture AFTER metrics
            print("📈 Step 7: Capturing AFTER training metrics")
            after_metrics = self._capture_after_metrics(minion_id, collection_name, rag_config)
            self.metrics_collector.record_after_metrics(after_metrics)
            training_service.update_training_job_progress(job_id, STEP_PROGRESS * 8, current_step=7)
            
            # Step 8: Calculate improvements
            print("📊 Step 8: Calculating improvements")
            improvements = self._calculate_improvements(before_metrics, after_metrics)
            final_metrics = self.metrics_collector.calculate_improvements()
            training_service.update_training_job_progress(job_id, 100.0, current_step=8)
            
            # Calculate XP and level progression
            xp_gained = self._calculate_training_xp(refinement_stats, validation_results, test_results, improvements)
            
            # Save detailed training results
            training_results = self._save_training_results(
                job_id, minion_id, user_id,
                before_metrics, after_metrics, improvements,
                refinement_stats, validation_results, test_results,
                collection_name, rag_config
            )
            
            # Websockets disabled - training completed successfully
            
            print(f"\n{'='*60}")
            print(f"✅ RAG Training Completed Successfully!")
            print(f"   Minion ID: {minion_id}")
            print(f"   XP Gained: {xp_gained}")
            print(f"   Accuracy Improvement: {improvements.get('accuracy', 0):.1f}%")
            print(f"   Speed Improvement: {improvements.get('speed', 0):.1f}%")
            print(f"   Knowledge Improvement: {improvements.get('knowledge', 0):.1f}%")
            print(f"{'='*60}\n")
            
            return {
                'success': True,
                'minion_id': minion_id,
                'xp_gained': xp_gained,
                'before_metrics': before_metrics,
                'after_metrics': after_metrics,
                'improvements': improvements,
                'refinement_stats': refinement_stats,
                'validation_results': validation_results,
                'test_results': test_results,
                'final_metrics': final_metrics,
                'collection_name': collection_name,
                'training_results_id': training_results.get('id')
            }
            
        except Exception as e:
            # Websockets disabled - training failed
            
            print(f"❌ RAG Training Failed: {str(e)}")
            return {
                'success': False,
                'error': str(e),
                'minion_id': minion_id
            }
    
    def _extract_text_from_file(self, file_path: str) -> Optional[str]:
        """Extract text from uploaded file (PDF, DOC, DOCX, TXT, MD)"""
        if not os.path.exists(file_path):
            print(f"   ⚠️ File not found: {file_path}")
            return None
        
        file_ext = os.path.splitext(file_path)[1].lower()
        
        try:
            if file_ext == '.txt' or file_ext == '.md':
                # Plain text files - try multiple encodings
                encodings = ['utf-8', 'latin-1', 'cp1252', 'iso-8859-1']
                for encoding in encodings:
                    try:
                        with open(file_path, 'r', encoding=encoding) as f:
                            return f.read()
                    except UnicodeDecodeError:
                        continue
                # If all encodings fail, try with error handling
                with open(file_path, 'r', encoding='utf-8', errors='replace') as f:
                    return f.read()
            
            elif file_ext == '.docx' and DOCX_AVAILABLE:
                # DOCX files (modern Word format)
                doc = docx.Document(file_path)
                paragraphs = [para.text for para in doc.paragraphs]
                return '\n'.join(paragraphs)
            
            elif file_ext == '.doc':
                # DOC files (older Word format) - requires textract or antiword
                # Note: .doc format is binary and harder to parse
                # For now, we'll try to use python-docx with error handling
                # or suggest converting to .docx
                print(f"   ⚠️ .doc format detected - attempting conversion...")
                try:
                    # Try using python-docx anyway (might work for some .doc files)
                    if DOCX_AVAILABLE:
                        doc = docx.Document(file_path)
                        paragraphs = [para.text for para in doc.paragraphs]
                        return '\n'.join(paragraphs)
                except Exception:
                    pass
                
                # If python-docx fails, suggest manual conversion
                print(f"   ⚠️ .doc files require conversion to .docx")
                print(f"   💡 Tip: Open the .doc file in Word and save as .docx")
                return None
            
            elif file_ext == '.pdf' and PDF_AVAILABLE:
                # PDF files
                with open(file_path, 'rb') as f:
                    if PDF_LIB == 'pypdf':
                        import pypdf
                        pdf_reader = pypdf.PdfReader(f)
                    else:
                        import PyPDF2
                        pdf_reader = PyPDF2.PdfReader(f)
                    
                    text = ''
                    for page in pdf_reader.pages:
                        text += page.extract_text() + '\n'
                    return text
            
            else:
                print(f"   ⚠️ Unsupported file type: {file_ext}")
                print(f"   💡 Supported formats: .txt, .md, .pdf, .docx, .doc")
                return None
                
        except Exception as e:
            print(f"   ❌ Error extracting text from {file_path}: {e}")
            import traceback
            traceback.print_exc()
            return None
    
    def _process_uploaded_files(self, rag_config: Dict[str, Any]) -> List[Dict[str, Any]]:
        """Process uploaded files and convert to dataset format"""
        uploaded_file_paths = rag_config.get('uploaded_file_paths', [])
        
        if not uploaded_file_paths:
            return []
        
        print(f"   📄 Processing {len(uploaded_file_paths)} uploaded file(s)...")
        
        file_datasets = []
        chunk_size = rag_config.get('chunkSize', 1000)
        chunk_overlap = rag_config.get('chunkOverlap', 100)
        
        for file_path in uploaded_file_paths:
            if not file_path:
                continue
            
            # Extract text from file
            text_content = self._extract_text_from_file(file_path)
            
            if text_content:
                file_name = os.path.basename(file_path)
                
                # Always chunk large files for better RAG retrieval
                # Use chunk_size instead of max_text_length for chunking decision
                if len(text_content) > chunk_size:
                    print(f"   📦 Chunking large file: {file_name} ({len(text_content)} chars → chunks)")
                    
                    # Chunk the content
                    chunks = self._chunk_text(text_content, chunk_size, chunk_overlap)
                    
                    for i, chunk in enumerate(chunks):
                        file_datasets.append({
                            'instruction': f"Content from {file_name} (chunk {i+1}/{len(chunks)})",
                            'output': chunk,
                            'input': '',
                            'source': 'uploaded_file',
                            'file_path': file_path,
                            'file_name': file_name,
                            'chunk_index': i,
                            'total_chunks': len(chunks)
                        })
                    
                    print(f"   ✓ Chunked into {len(chunks)} pieces")
                else:
                    # Small file - add as single document
                    file_datasets.append({
                        'instruction': f"Content from {file_name}",
                        'output': text_content,
                        'input': '',
                        'source': 'uploaded_file',
                        'file_path': file_path,
                        'file_name': file_name
                    })
                
                print(f"   ✓ Processed: {file_name} ({len(text_content)} chars)")
            else:
                print(f"   ⚠️ Skipped: {os.path.basename(file_path)}")
        
        print(f"   ✅ Processed {len(file_datasets)} document(s) from {len(uploaded_file_paths)} file(s)")
        return file_datasets
    
    def _chunk_text(self, text: str, chunk_size: int, chunk_overlap: int) -> List[str]:
        """Split text into overlapping chunks"""
        if len(text) <= chunk_size:
            return [text]
        
        chunks = []
        start = 0
        
        while start < len(text):
            end = start + chunk_size
            chunk = text[start:end]
            chunks.append(chunk)
            
            # Move start position forward by chunk_size - overlap
            start += chunk_size - chunk_overlap
            
            # Prevent infinite loop
            if start >= len(text):
                break
        
        return chunks
    
    def _refine_datasets(self, datasets: List[Dict[str, Any]], rag_config: Dict[str, Any]) -> Tuple[List[Dict[str, Any]], Dict[str, Any]]:
        """Refine datasets using the dataset refiner"""
        print("   🧹 Refining datasets...")
        
        # First, process uploaded files if any
        uploaded_file_datasets = self._process_uploaded_files(rag_config)
        
        # Combine all datasets
        combined_data = []
        
        # Add uploaded file datasets
        combined_data.extend(uploaded_file_datasets)
        
        # Add database datasets
        for dataset in datasets:
            if isinstance(dataset, dict) and 'all_samples' in dataset:
                combined_data.extend(dataset['all_samples'])
            elif isinstance(dataset, dict) and 'data' in dataset:
                combined_data.extend(dataset['data'])
            elif isinstance(dataset, list):
                combined_data.extend(dataset)
        
        if not combined_data:
            print("   ⚠️ No datasets to refine (neither uploaded files nor database datasets)")
            return [], {
                'original_count': 0,
                'refined_count': 0,
                'removed_empty': 0,
                'removed_duplicates': 0,
                'removed_length_filter': 0,
                'removed_malformed': 0
            }
        
        # Separate uploaded files from dataset items
        # Uploaded files are already chunked, so we don't need length filtering for them
        uploaded_items = [item for item in combined_data if item.get('source') == 'uploaded_file']
        dataset_items = [item for item in combined_data if item.get('source') != 'uploaded_file']
        
        # Refine dataset items (with length filtering)
        refined_dataset_items = []
        if dataset_items:
            # Use a much higher limit for datasets (or None to disable)
            # Uploaded files are already chunked, so they bypass this entirely
            max_dataset_length = rag_config.get('max_text_length', 100000)  # Increased from 10,000 to 100,000
            
            refined_dataset_items, stats = self.dataset_refiner.refine_dataset(
                dataset_items,
                min_text_length=rag_config.get('min_text_length', 10),
                max_text_length=max_dataset_length
            )
        
        # Uploaded files are already chunked and validated, so just add them directly
        # (they've already been processed and chunked in _process_uploaded_files)
        final_refined_data = uploaded_items + refined_dataset_items
        
        print(f"   ✓ Refined {len(dataset_items)} dataset items → {len(refined_dataset_items)} items")
        if uploaded_items:
            print(f"   ✓ Included {len(uploaded_items)} uploaded file chunk(s) (already processed)")
        
        return final_refined_data, stats if dataset_items else {
            'original_count': len(uploaded_items),
            'refined_count': len(uploaded_items),
            'removed_empty': 0,
            'removed_duplicates': 0,
            'removed_length_filter': 0,
            'removed_malformed': 0
        }
    
    def _create_knowledge_base(self, job_id: int, datasets: List[Dict[str, Any]], rag_config: Dict[str, Any]) -> str:
        """Create ChromaDB knowledge base using rag_config parameters"""
        # Use collection name from rag_config, fallback to generated name
        collection_name = rag_config.get('collectionName', f"rag_training_{job_id}_{int(time.time())}")
        collection_description = rag_config.get('collectionDescription', f"RAG training collection for job {job_id}")
        
        print(f"   📚 Creating collection: {collection_name}")
        print(f"   📝 Description: {collection_description}")
        
        # Apply knowledge base strategy
        knowledge_base_strategy = rag_config.get('knowledgeBaseStrategy', 'create_new')
        print(f"   🎯 Knowledge Base Strategy: {knowledge_base_strategy}")
        
        if knowledge_base_strategy == 'use_existing':
            existing_collection = rag_config.get('existingCollection', '')
            if existing_collection:
                print(f"   ♻️ Using existing collection: {existing_collection}")
                return existing_collection
            else:
                print(f"   ⚠️ No existing collection specified, creating new one")
        
        # Create collection with config parameters
        success = self.chromadb_service.create_collection(
            collection_name, 
            collection_description
        )
        
        if not success:
            raise Exception(f"Failed to create collection: {collection_name}")
        
        # Apply update strategy
        update_strategy = rag_config.get('updateStrategy', 'append')
        print(f"   🔄 Update Strategy: {update_strategy}")
        
        # Apply chunking parameters
        chunk_size = rag_config.get('chunkSize', 1000)
        chunk_overlap = rag_config.get('chunkOverlap', 100)
        print(f"   📏 Chunking: size={chunk_size}, overlap={chunk_overlap}")
        
        # Ingest datasets (ChromaDB handles chunking internally)
        print(f"   📥 Ingesting {len(datasets)} documents...")
        ingest_success = self.chromadb_service.ingest_dataset(
            collection_name, 
            datasets
        )
        
        if not ingest_success:
            raise Exception(f"Failed to ingest datasets into collection: {collection_name}")
        
        print(f"   ✓ Knowledge base created successfully")
        return collection_name
    
    def _create_embeddings(self, collection_name: str, datasets: List[Dict[str, Any]], rag_config: Dict[str, Any]) -> Dict[str, Any]:
        """Create embeddings for the knowledge base"""
        print("   🔮 Creating embeddings...")
        
        # ChromaDB automatically creates embeddings during ingestion
        # This step is mainly for validation and statistics
        
        # Test embedding creation
        test_query = "test query for embedding validation"
        results = self.chromadb_service.query_collection(collection_name, test_query, 1)
        
        if not results:
            raise Exception("Failed to create embeddings - no results from test query")
        
        print(f"   ✓ Embeddings created successfully")
        return {
            'total_embeddings': len(datasets),
            'embedding_model': rag_config.get('embeddingModel', 'all-MiniLM-L6-v2'),
            'test_query_successful': True
        }
    
    def _update_minion_with_rag(self, minion_id: int, collection_name: str, rag_config: Dict[str, Any]) -> Dict[str, Any]:
        """Update minion configuration with RAG settings"""
        print(f"   🤖 Updating minion {minion_id} with RAG configuration...")
        
        try:
            # Update minion in database with RAG configuration
            from app.services.minion_service import MinionService
            service = MinionService()
            
            # Prepare RAG configuration data
            rag_update_data = {
                'rag_enabled': True,
                'rag_collection_name': collection_name,
                'top_k': rag_config.get('topK', 4),
                'similarity_threshold': rag_config.get('similarityThreshold', 0.7),
                'retrieval_method': rag_config.get('retrievalMethod', 'semantic'),
                'enable_contextual_compression': rag_config.get('enableContextualCompression', False),
                'enable_source_citation': rag_config.get('enableSourceCitation', False),
                'enable_query_expansion': rag_config.get('enableQueryExpansion', False),
                'embedding_model': rag_config.get('embeddingModel', 'all-MiniLM-L6-v2'),
                'chunk_size': rag_config.get('chunkSize', 1000),
                'chunk_overlap': rag_config.get('chunkOverlap', 100)
            }
            
            # Update minion with RAG configuration
            result = service.update_minion(minion_id, rag_update_data)
            
            if not result.get('success', False):
                raise Exception(f"Failed to update minion: {result.get('error', 'Unknown error')}")
            
            print(f"   ✓ Minion updated with RAG configuration:")
            print(f"      Collection: {collection_name}")
            print(f"      TopK: {rag_update_data['top_k']}")
            print(f"      Similarity Threshold: {rag_update_data['similarity_threshold']}")
            print(f"      Retrieval Method: {rag_update_data['retrieval_method']}")
            print(f"      Contextual Compression: {rag_update_data['enable_contextual_compression']}")
            print(f"      Source Citation: {rag_update_data['enable_source_citation']}")
            
            return {
                'minion_id': minion_id,
                'collection_name': collection_name,
                'rag_config_applied': True,
                'updated_at': datetime.now().isoformat(),
                'rag_config': rag_update_data
            }
            
        except Exception as e:
            print(f"   ❌ Failed to update minion with RAG config: {e}")
            return {
                'minion_id': minion_id,
                'collection_name': collection_name,
                'rag_config_applied': False,
                'error': str(e),
                'updated_at': datetime.now().isoformat()
            }
    
    def _validate_training(self, minion_id: int, collection_name: str, rag_config: Dict[str, Any]) -> Dict[str, Any]:
        """Validate the training results"""
        print("   ✅ Validating training results...")
        
        # Test queries to validate the knowledge base
        test_queries = [
            "What is the main topic?",
            "Can you explain this concept?",
            "How does this work?"
        ]
        
        validation_results = {
            'tests_total': len(test_queries),
            'tests_passed': 0,
            'overall_score': 0,
            'query_results': []
        }
        
        for query in test_queries:
            try:
                # Apply retrieval parameters from rag_config
                top_k = rag_config.get('topK', 4)
                similarity_threshold = rag_config.get('similarityThreshold', 0.7)
                retrieval_method = rag_config.get('retrievalMethod', 'semantic')
                
                print(f"   🔍 Retrieval: topK={top_k}, threshold={similarity_threshold}, method={retrieval_method}")
                
                results = self.chromadb_service.query_collection(
                    collection_name, 
                    query, 
                    n_results=top_k
                )
                
                if results and len(results) > 0:
                    validation_results['tests_passed'] += 1
                    validation_results['query_results'].append({
                        'query': query,
                        'results_count': len(results),
                        'success': True
                    })
                else:
                    validation_results['query_results'].append({
                        'query': query,
                        'results_count': 0,
                        'success': False
                    })
                    
            except Exception as e:
                validation_results['query_results'].append({
                    'query': query,
                    'error': str(e),
                    'success': False
                })
        
        # Calculate overall score
        validation_results['overall_score'] = (
            validation_results['tests_passed'] / validation_results['tests_total']
        ) * 100
        
        print(f"   ✓ Validation completed: {validation_results['tests_passed']}/{validation_results['tests_total']} tests passed")
        return validation_results
    
    def _test_minion_performance(self, minion_id: int, collection_name: str, rag_config: Dict[str, Any]) -> Dict[str, Any]:
        """Test minion performance with RAG"""
        print("   🧪 Testing minion performance...")
        
        # Test queries for performance evaluation
        performance_tests = [
            {
                'query': 'What can you help me with?',
                'expected_keywords': ['help', 'assist', 'support']
            },
            {
                'query': 'Explain the main concepts',
                'expected_keywords': ['concept', 'explain', 'main']
            }
        ]
        
        test_results = {
            'tests_total': len(performance_tests),
            'tests_passed': 0,
            'performance_score': 0,
            'test_details': []
        }
        
        for test in performance_tests:
            try:
                # Apply retrieval parameters from rag_config
                top_k = rag_config.get('topK', 4)
                similarity_threshold = rag_config.get('similarityThreshold', 0.7)
                retrieval_method = rag_config.get('retrievalMethod', 'semantic')
                
                # Query the knowledge base with config parameters
                kb_results = self.chromadb_service.query_collection(
                    collection_name,
                    test['query'],
                    n_results=top_k
                )
                
                # Check if results contain expected keywords
                if kb_results:
                    result_text = ' '.join([r.get('document', '') for r in kb_results])
                    keywords_found = sum(1 for keyword in test['expected_keywords'] 
                                       if keyword.lower() in result_text.lower())
                    
                    test_passed = keywords_found > 0
                    if test_passed:
                        test_results['tests_passed'] += 1
                    
                    test_results['test_details'].append({
                        'query': test['query'],
                        'keywords_found': keywords_found,
                        'passed': test_passed,
                        'results_count': len(kb_results)
                    })
                else:
                    test_results['test_details'].append({
                        'query': test['query'],
                        'keywords_found': 0,
                        'passed': False,
                        'results_count': 0
                    })
                    
            except Exception as e:
                test_results['test_details'].append({
                    'query': test['query'],
                    'error': str(e),
                    'passed': False
                })
        
        # Calculate performance score
        test_results['performance_score'] = (
            test_results['tests_passed'] / test_results['tests_total']
        ) * 100
        
        print(f"   ✓ Performance testing completed: {test_results['tests_passed']}/{test_results['tests_total']} tests passed")
        return test_results
    
    def _capture_before_metrics(self, minion_id: int, minion_config: Dict[str, Any]) -> Dict[str, Any]:
        """Capture minion performance metrics before training"""
        print("   📊 Capturing BEFORE training metrics...")
        
        # Test queries to measure baseline performance
        baseline_queries = [
            "What can you help me with?",
            "Explain a concept from your knowledge",
            "How do you work?",
            "What is your expertise?"
        ]
        
        before_metrics = {
            'timestamp': datetime.now().isoformat(),
            'minion_id': minion_id,
            'baseline_tests': [],
            'response_times': [],
            'accuracy_scores': [],
            'knowledge_coverage': 0,
            'response_quality': 0,
            'overall_score': 0
        }
        
        total_response_time = 0
        total_accuracy = 0
        
        for query in baseline_queries:
            try:
                start_time = time.time()
                
                # Test minion without RAG (baseline)
                response = self._test_minion_response(minion_id, query, use_rag=False)
                
                response_time = time.time() - start_time
                total_response_time += response_time
                
                # Use THE_ANSWER grading system for comprehensive evaluation
                kb_results = []  # No RAG before training
                grading_results = self._implement_the_answer_grading_system(query, response, kb_results, minion_config)
                accuracy_score = grading_results['task_accuracy']
                total_accuracy += accuracy_score
                
                before_metrics['baseline_tests'].append({
                    'query': query,
                    'response': response,
                    'response_time': response_time,
                    'accuracy_score': accuracy_score,
                    'success': len(response) > 0
                })
                
                before_metrics['response_times'].append(response_time)
                before_metrics['accuracy_scores'].append(accuracy_score)
                
            except Exception as e:
                before_metrics['baseline_tests'].append({
                    'query': query,
                    'error': str(e),
                    'response_time': 0,
                    'accuracy_score': 0,
                    'success': False
                })
        
        # Calculate averages
        if before_metrics['response_times']:
            before_metrics['avg_response_time'] = sum(before_metrics['response_times']) / len(before_metrics['response_times'])
        else:
            before_metrics['avg_response_time'] = 0
            
        if before_metrics['accuracy_scores']:
            before_metrics['avg_accuracy'] = sum(before_metrics['accuracy_scores']) / len(before_metrics['accuracy_scores'])
        else:
            before_metrics['avg_accuracy'] = 0
        
        # Calculate overall score
        before_metrics['overall_score'] = (
            before_metrics['avg_accuracy'] * 0.7 +  # 70% accuracy
            (100 - min(before_metrics['avg_response_time'] * 10, 100)) * 0.3  # 30% speed (inverse of response time)
        )
        
        print(f"   ✓ BEFORE metrics captured:")
        print(f"      Avg Response Time: {before_metrics['avg_response_time']:.2f}s")
        print(f"      Avg Accuracy: {before_metrics['avg_accuracy']:.1f}%")
        print(f"      Overall Score: {before_metrics['overall_score']:.1f}%")
        
        return before_metrics
    
    def _capture_after_metrics(self, minion_id: int, collection_name: str, rag_config: Dict[str, Any]) -> Dict[str, Any]:
        """Capture minion performance metrics after training"""
        print("   📈 Capturing AFTER training metrics...")
        
        # Same test queries to measure improvement
        test_queries = [
            "What can you help me with?",
            "Explain a concept from your knowledge",
            "How do you work?",
            "What is your expertise?"
        ]
        
        after_metrics = {
            'timestamp': datetime.now().isoformat(),
            'minion_id': minion_id,
            'collection_name': collection_name,
            'rag_enabled': True,
            'performance_tests': [],
            'response_times': [],
            'accuracy_scores': [],
            'knowledge_coverage': 0,
            'response_quality': 0,
            'overall_score': 0
        }
        
        total_response_time = 0
        total_accuracy = 0
        knowledge_coverage_score = 0
        
        for query in test_queries:
            try:
                start_time = time.time()
                
                # Test minion with RAG enabled
                response = self._test_minion_response(minion_id, query, use_rag=True, collection_name=collection_name)
                
                response_time = time.time() - start_time
                total_response_time += response_time
                
                # Check knowledge base coverage first
                kb_results = self.chromadb_service.query_collection(collection_name, query, 3)
                
                # Use THE_ANSWER grading system for comprehensive evaluation
                grading_results = self._implement_the_answer_grading_system(query, response, kb_results, minion_config)
                accuracy_score = grading_results['task_accuracy']
                total_accuracy += accuracy_score
                
                # Knowledge utilization score from THE_ANSWER system
                coverage_score = self._evaluate_knowledge_coverage(query, response, kb_results, minion_config)
                knowledge_coverage_score += coverage_score
                
                after_metrics['performance_tests'].append({
                    'query': query,
                    'response': response,
                    'response_time': response_time,
                    'accuracy_score': accuracy_score,
                    'knowledge_coverage': coverage_score,
                    'kb_results_count': len(kb_results),
                    'success': len(response) > 0
                })
                
                after_metrics['response_times'].append(response_time)
                after_metrics['accuracy_scores'].append(accuracy_score)
                
            except Exception as e:
                after_metrics['performance_tests'].append({
                    'query': query,
                    'error': str(e),
                    'response_time': 0,
                    'accuracy_score': 0,
                    'knowledge_coverage': 0,
                    'success': False
                })
        
        # Calculate averages
        if after_metrics['response_times']:
            after_metrics['avg_response_time'] = sum(after_metrics['response_times']) / len(after_metrics['response_times'])
        else:
            after_metrics['avg_response_time'] = 0
            
        if after_metrics['accuracy_scores']:
            after_metrics['avg_accuracy'] = sum(after_metrics['accuracy_scores']) / len(after_metrics['accuracy_scores'])
        else:
            after_metrics['avg_accuracy'] = 0
            
        after_metrics['knowledge_coverage'] = knowledge_coverage_score / len(test_queries)
        
        # Calculate overall score
        after_metrics['overall_score'] = (
            after_metrics['avg_accuracy'] * 0.5 +  # 50% accuracy
            after_metrics['knowledge_coverage'] * 0.3 +  # 30% knowledge coverage
            (100 - min(after_metrics['avg_response_time'] * 10, 100)) * 0.2  # 20% speed
        )
        
        print(f"   ✓ AFTER metrics captured:")
        print(f"      Avg Response Time: {after_metrics['avg_response_time']:.2f}s")
        print(f"      Avg Accuracy: {after_metrics['avg_accuracy']:.1f}%")
        print(f"      Knowledge Coverage: {after_metrics['knowledge_coverage']:.1f}%")
        print(f"      Overall Score: {after_metrics['overall_score']:.1f}%")
        
        return after_metrics
    
    def _calculate_improvements(self, before_metrics: Dict[str, Any], after_metrics: Dict[str, Any]) -> Dict[str, Any]:
        """Calculate improvement percentages using real before/after metrics"""
        print("   📊 Calculating improvements from real metrics...")
        
        improvements = {
            'accuracy': 0,
            'speed': 0,
            'knowledge': 0,
            'overall': 0,
            'response_time_improvement': 0,
            'knowledge_coverage_gain': 0
        }
        
        # Accuracy improvement (real before/after comparison)
        before_accuracy = before_metrics.get('avg_accuracy', 0)
        after_accuracy = after_metrics.get('avg_accuracy', 0)
        if before_accuracy > 0:
            improvements['accuracy'] = ((after_accuracy - before_accuracy) / before_accuracy) * 100
        elif after_accuracy > 0:
            # If no before accuracy, knowledge is new capability
            improvements['accuracy'] = after_accuracy
        
        # Speed improvement (response time reduction - real measurement)
        before_speed = before_metrics.get('avg_response_time', 0)
        after_speed = after_metrics.get('avg_response_time', 0)
        if before_speed > 0 and after_speed > 0:
            improvements['speed'] = ((before_speed - after_speed) / before_speed) * 100
            improvements['response_time_improvement'] = improvements['speed']
        elif after_speed > 0:
            # If no before speed, this is new capability
            improvements['speed'] = 0  # Can't measure improvement without baseline
        
        # Knowledge improvement (real knowledge base utilization)
        before_knowledge = before_metrics.get('knowledge_coverage', 0)
        after_knowledge = after_metrics.get('knowledge_coverage', 0)
        if before_knowledge > 0:
            improvements['knowledge'] = ((after_knowledge - before_knowledge) / before_knowledge) * 100
        elif after_knowledge > 0:
            # If no before knowledge, this is new capability
            improvements['knowledge'] = after_knowledge
        improvements['knowledge_coverage_gain'] = improvements['knowledge']
        
        # Overall improvement (weighted average of all metrics)
        accuracy_weight = 0.4
        speed_weight = 0.2
        knowledge_weight = 0.4
        
        overall_score = (
            (improvements['accuracy'] * accuracy_weight) +
            (improvements['speed'] * speed_weight) +
            (improvements['knowledge'] * knowledge_weight)
        )
        improvements['overall'] = overall_score
        
        print(f"   ✓ Real improvements calculated:")
        print(f"      Before Accuracy: {before_accuracy:.1f}% → After: {after_accuracy:.1f}%")
        print(f"      Before Speed: {before_speed:.2f}s → After: {after_speed:.2f}s")
        print(f"      Before Knowledge: {before_knowledge:.1f}% → After: {after_knowledge:.1f}%")
        print(f"      Accuracy: {improvements['accuracy']:+.1f}%")
        print(f"      Speed: {improvements['speed']:+.1f}%")
        print(f"      Knowledge: +{improvements['knowledge']:.1f}%")
        print(f"      Overall: {improvements['overall']:+.1f}%")
        
        return improvements
    
    def _test_minion_response(self, minion_id: int, query: str, use_rag: bool = False, collection_name: str = None) -> str:
        """Test minion response with or without RAG using ExternalAPIService"""
        try:
            # Fetch minion config from DB to obtain provider/model/api_key/base_url
            from app.services.minion_service import MinionService
            service = MinionService()
            minion = service.get_minion_by_id(minion_id)
            if not minion:
                raise Exception(f"Minion {minion_id} not found")

            model_config = {
                'provider': minion.get('provider'),
                'model_id': minion.get('model_id'),
                'api_key': minion.get('api_key'),
                'base_url': minion.get('base_url'),
                'temperature': minion.get('temperature', 0.7),
                'top_p': minion.get('top_p', 0.9),
                'max_tokens': minion.get('max_tokens', 1024),
                'system_prompt': minion.get('system_prompt', '')
            }

            # If RAG is requested, include retrieval context from ChromaDB
            prompt = query
            if use_rag and collection_name:
                # Apply retrieval parameters from minion config or defaults
                top_k = minion.get('top_k', 3)  # Default to 3 for testing
                similarity_threshold = minion.get('similarity_threshold', 0.7)
                retrieval_method = minion.get('retrieval_method', 'semantic')
                
                kb_results = self.chromadb_service.query_collection(
                    collection_name, 
                    query, 
                    n_results=top_k
                )
                if kb_results:
                    # Apply contextual compression if enabled
                    enable_contextual_compression = minion.get('enable_contextual_compression', False)
                    if enable_contextual_compression:
                        # Compress context to most relevant parts
                        kb_context = self._compress_context(kb_results, query)
                    else:
                        kb_context = "\n\n".join([r.get('document', '') for r in kb_results])
                    
                    # Apply source citation if enabled
                    enable_source_citation = minion.get('enable_source_citation', False)
                    if enable_source_citation:
                        prompt = f"<KnowledgeBase>\n{kb_context}\n</KnowledgeBase>\n\n{query}\n\nPlease cite sources when using information from the knowledge base."
                    else:
                        prompt = f"<KnowledgeBase>\n{kb_context}\n</KnowledgeBase>\n\n{query}"

            # Call external API (stream then join)
            response_chunks = []
            for chunk in self.external_api_service.call_model(model_config, prompt, max_tokens=512):
                response_chunks.append(chunk)

            response_text = ''.join(response_chunks)
            return response_text

        except Exception as e:
            print(f"❌ Error calling minion API for minion {minion_id}: {e}")
            return f"Error: {str(e)}"
    
    def _evaluate_response_quality(self, query: str, response: str, minion_config: Dict[str, Any] = None) -> float:
        """Evaluate response quality using THE_ANSWER grading system (0-100)"""
        if not response or len(response) < 10:
            return 0
        
        try:
            # Use THE_ANSWER grading system - Task Accuracy evaluation
            evaluation_prompt = f"""
            You are evaluating an AI assistant's task completion accuracy. Rate 0-100 based on:
            
            Task: "{query}"
            Response: "{response}"
            
            Grading Criteria (THE_ANSWER System):
            - Task Accuracy (0-100): Did the response correctly complete the requested task?
            - Completeness (0-100): Was the task fully addressed without missing key components?
            - Correctness (0-100): Is the information provided accurate and reliable?
            
            Calculate the average of these three scores and respond with only the final number (0-100).
            """
            
            # Use minion's own API for evaluation (as per architecture)
            eval_config = None
            
            if minion_config and minion_config.get('api_key'):
                # Use minion's own API key for evaluation
                eval_config = {
                    'provider': minion_config.get('provider', 'nvidia'),
                    'model_id': minion_config.get('model_id', 'nvidia/llama-3.3-nemotron-super-49b-v1.5'),
                    'api_key': minion_config.get('api_key'),
                    'base_url': minion_config.get('base_url', 'https://integrate.api.nvidia.com/v1'),
                    'temperature': 0.1,  # Low temperature for consistent evaluation
                    'max_tokens': 10
                }
                print(f"   🔍 Using minion's API ({eval_config['provider']}) for task accuracy evaluation")
            else:
                # Fallback to OpenAI API key
                openai_key = os.getenv('OPENAI_API_KEY')
                if not openai_key:
                    print(f"   ⚠️ No API key available, using fallback evaluation")
                    return self._fallback_response_evaluation(query, response)
                
                eval_config = {
                    'provider': 'openai',
                    'model_id': 'gpt-3.5-turbo',
                    'api_key': openai_key,
                    'base_url': 'https://api.openai.com/v1',
                    'temperature': 0.1,
                    'max_tokens': 10
                }
                print(f"   🔍 Using OpenAI API key for evaluation")
            
            # Call evaluation API with realistic delay for real training
            import time
            time.sleep(2)  # Realistic delay for AI evaluation
            
            eval_response = ""
            for chunk in self.external_api_service.call_model(eval_config, evaluation_prompt, max_tokens=10):
                eval_response += chunk
            
            # Extract numeric score
            import re
            score_match = re.search(r'\b(\d{1,3})\b', eval_response.strip())
            if score_match:
                score = int(score_match.group(1))
                print(f"   ✅ Task accuracy score: {score}")
                return min(max(score, 0), 100)  # Clamp between 0-100
            else:
                print(f"   ⚠️ Could not parse evaluation response: {eval_response}")
                return self._fallback_response_evaluation(query, response)
                
        except Exception as e:
            print(f"   ⚠️ Task accuracy evaluation failed: {e}, using fallback")
            return self._fallback_response_evaluation(query, response)
    
    def _fallback_response_evaluation(self, query: str, response: str) -> float:
        """Fallback evaluation when AI evaluation fails"""
        if not response or len(response) < 10:
            return 0
        
        score = 30  # Base score (lower than before)
        
        # Length bonus (more conservative)
        if len(response) > 50:
            score += 15
        if len(response) > 100:
            score += 10
        
        # Relevance bonus (improved keyword matching)
        query_words = set(query.lower().split())
        response_words = set(response.lower().split())
        overlap = len(query_words.intersection(response_words))
        if overlap > 0:
            score += min(overlap * 3, 15)  # More conservative bonus
        
        # Penalty for very short responses
        if len(response) < 20:
            score -= 20
        
        return min(max(score, 0), 100)  # Clamp between 0-100
    
    def _evaluate_knowledge_coverage(self, query: str, response: str, kb_results: List[Dict], minion_config: Dict[str, Any] = None) -> float:
        """Evaluate knowledge utilization using THE_ANSWER grading system (0-100)"""
        if not kb_results or not response:
            return 0
        
        try:
            # Extract knowledge base content
            kb_content = "\n".join([r.get('document', '') for r in kb_results])
            
            # Use THE_ANSWER grading system - Knowledge Utilization evaluation
            evaluation_prompt = f"""
            You are evaluating how well an AI assistant utilizes provided knowledge. Rate 0-100 based on:
            
            Task: "{query}"
            Available Knowledge: "{kb_content[:1000]}..."  # Truncated for token limits
            AI Response: "{response}"
            
            Grading Criteria (THE_ANSWER System):
            - Knowledge Utilization (0-100): How much relevant knowledge was effectively used?
            - Information Integration (0-100): How well was knowledge integrated into the response?
            - Knowledge Accuracy (0-100): Was the knowledge applied correctly and accurately?
            
            Calculate the average of these three scores and respond with only the final number (0-100).
            """
            
            # Use minion's own API for evaluation (as per architecture)
            eval_config = None
            
            if minion_config and minion_config.get('api_key'):
                # Use minion's own API key for evaluation
                eval_config = {
                    'provider': minion_config.get('provider', 'nvidia'),
                    'model_id': minion_config.get('model_id', 'nvidia/llama-3.3-nemotron-super-49b-v1.5'),
                    'api_key': minion_config.get('api_key'),
                    'base_url': minion_config.get('base_url', 'https://integrate.api.nvidia.com/v1'),
                    'temperature': 0.1,
                    'max_tokens': 10
                }
                print(f"   🔍 Using minion's API ({eval_config['provider']}) for knowledge utilization evaluation")
            else:
                # Fallback to OpenAI API key
                openai_key = os.getenv('OPENAI_API_KEY')
                if not openai_key:
                    print(f"   ⚠️ No API key available, using fallback knowledge evaluation")
                    return self._fallback_knowledge_coverage(query, response, kb_results)
                
                eval_config = {
                    'provider': 'openai',
                    'model_id': 'gpt-3.5-turbo',
                    'api_key': openai_key,
                    'base_url': 'https://api.openai.com/v1',
                    'temperature': 0.1,
                    'max_tokens': 10
                }
            
            # Call evaluation API with realistic delay for real training
            import time
            time.sleep(2)  # Realistic delay for AI evaluation
            
            eval_response = ""
            for chunk in self.external_api_service.call_model(eval_config, evaluation_prompt, max_tokens=10):
                eval_response += chunk
            
            # Extract numeric score
            import re
            score_match = re.search(r'\b(\d{1,3})\b', eval_response.strip())
            if score_match:
                score = int(score_match.group(1))
                print(f"   ✅ Knowledge utilization score: {score}")
                return min(max(score, 0), 100)
            else:
                print(f"   ⚠️ Could not parse knowledge evaluation response: {eval_response}")
                return self._fallback_knowledge_coverage(query, response, kb_results)
                
        except Exception as e:
            print(f"   ⚠️ Knowledge utilization evaluation failed: {e}, using fallback")
            return self._fallback_knowledge_coverage(query, response, kb_results)
    
    def _fallback_knowledge_coverage(self, query: str, response: str, kb_results: List[Dict]) -> float:
        """Fallback knowledge coverage evaluation"""
        if not kb_results:
            return 0
        
        # Simple content overlap analysis
        kb_content = " ".join([r.get('document', '') for r in kb_results])
        kb_words = set(kb_content.lower().split())
        response_words = set(response.lower().split())
        
        # Calculate overlap percentage
        overlap = len(kb_words.intersection(response_words))
        total_kb_words = len(kb_words)
        
        if total_kb_words == 0:
            return 0
        
        coverage_percentage = (overlap / total_kb_words) * 100
        return min(coverage_percentage, 100)
    
    def _save_training_results(self, job_id: int, minion_id: int, user_id: int,
                             before_metrics: Dict[str, Any], after_metrics: Dict[str, Any],
                             improvements: Dict[str, Any], refinement_stats: Dict[str, Any],
                             validation_results: Dict[str, Any], test_results: Dict[str, Any],
                             collection_name: str, rag_config: Dict[str, Any]) -> Dict[str, Any]:
        """Save detailed training results to database"""
        print("   💾 Saving training results to DB...")

        # Persist using SQLAlchemy TrainingResult model
        from model.training_results import TrainingResult
        from database.session import get_session_sync

        session = get_session_sync()
        try:
            tr = TrainingResult(
                job_id=job_id,
                minion_id=minion_id,
                user_id=user_id,
                training_type='RAG',
                collection_name=collection_name,
                before_metrics=before_metrics,
                after_metrics=after_metrics,
                improvements=improvements,
                refinement_stats=refinement_stats,
                validation_results=validation_results,
                test_results=test_results,
                rag_config=rag_config,
                minion_config={},
                accuracy_improvement=improvements.get('accuracy', 0),
                speed_improvement=improvements.get('speed', 0),
                knowledge_improvement=improvements.get('knowledge', 0),
                overall_improvement=improvements.get('overall', 0),
                xp_gained=self._calculate_training_xp(refinement_stats, validation_results, test_results, improvements)
            )

            session.add(tr)
            session.commit()
            session.refresh(tr)

            print(f"   ✓ Training results saved with DB id: {tr.id}")
            return tr.to_dict()
        except Exception as e:
            session.rollback()
            print(f"❌ Failed to save training results: {e}")
            return {
                'id': f"training_result_{job_id}_{int(time.time())}",
                'error': str(e)
            }
        finally:
            session.close()
    
    def _calculate_training_xp(self, refinement_stats: Dict[str, Any], validation_results: Dict[str, Any], test_results: Dict[str, Any], improvements: Dict[str, Any]) -> int:
        """Calculate XP gained from training using THE_ANSWER piecewise linear interpolation"""
        from bisect import bisect_right
        
        # Get refined items count (dataset size)
        refined_items = refinement_stats.get('final_count', 0)
        
        # THE_ANSWER XP breakpoints from architectural plan
        BREAKPOINTS = [
            (0, 0),
            (10, 20),
            (100, 200),
            (500, 750),
            (4800, 1400),  # Large dataset (CodeAlpaca)
            (20000, 2600)
        ]
        
        def dataset_lines_to_xp(lines: int) -> int:
            """Convert dataset size to XP using piecewise-linear interpolation"""
            if lines <= 0:
                return 0
            
            xs = [p[0] for p in BREAKPOINTS]
            ys = [p[1] for p in BREAKPOINTS]
            
            if lines >= xs[-1]:
                return ys[-1]
            
            idx = bisect_right(xs, lines) - 1
            x0, y0 = xs[idx], ys[idx]
            x1, y1 = xs[idx+1], ys[idx+1]
            
            # Linear interpolation
            t = (lines - x0) / (x1 - x0)
            xp = y0 + t * (y1 - y0)
            return int(round(xp))
        
        # 1. Base XP from dataset size (using THE_ANSWER interpolation)
        base_xp = dataset_lines_to_xp(refined_items)
        
        # 2. Quality bonus (10-20% of base)
        quality_score = refinement_stats.get('quality_score', 0)
        quality_bonus = int(base_xp * (quality_score / 100) * 0.2)
        
        # 3. Validation bonus (THE_ANSWER system)
        validation_score = validation_results.get('overall_score', 0)
        if validation_score == 100:
            validation_bonus = 200
        elif validation_score >= 90:
            validation_bonus = 150
        elif validation_score >= 80:
            validation_bonus = 100
        elif validation_score >= 70:
            validation_bonus = 50
        else:
            validation_bonus = 0
        
        # 4. Task completion bonus (THE_ANSWER system)
        task_accuracy = improvements.get('accuracy', 0)
        task_completion_bonus = int(max(0, task_accuracy) * 0.5)  # Up to 50 XP bonus
        
        total_xp = base_xp + quality_bonus + validation_bonus + task_completion_bonus
        
        print(f"   📊 THE_ANSWER XP Calculation:")
        print(f"      Dataset Size: {refined_items} items")
        print(f"      Base XP (interpolation): {base_xp}")
        print(f"      Quality Bonus ({quality_score}%): +{quality_bonus}")
        print(f"      Validation Bonus ({validation_score}%): +{validation_bonus}")
        print(f"      Task Completion Bonus ({task_accuracy:+.1f}%): +{task_completion_bonus}")
        print(f"      Total XP: {total_xp}")
        
        return total_xp
    
    def _implement_the_answer_grading_system(self, query: str, response: str, kb_results: List[Dict], minion_config: Dict[str, Any] = None) -> Dict[str, Any]:
        """Implement THE_ANSWER comprehensive grading system"""
        print("   📊 Implementing THE_ANSWER grading system...")
        
        grading_results = {
            'task_accuracy': 0,
            'revision_count': 0,  # Would track user revisions in real implementation
            'task_complexity_multiplier': 1.0,
            'skill_adherence': 0,
            'time_efficiency': 0,
            'overall_score': 0
        }
        
        # 1. Task Accuracy (0-100) - already implemented in _evaluate_response_quality
        grading_results['task_accuracy'] = self._evaluate_response_quality(query, response, minion_config)
        
        # 2. Task Complexity Multiplier (×1, ×2, ×3)
        complexity_keywords = {
            'simple': ['what', 'explain', 'define', 'list'],
            'medium': ['analyze', 'compare', 'evaluate', 'create'],
            'complex': ['design', 'implement', 'optimize', 'integrate', 'develop']
        }
        
        query_lower = query.lower()
        if any(word in query_lower for word in complexity_keywords['complex']):
            grading_results['task_complexity_multiplier'] = 3.0
        elif any(word in query_lower for word in complexity_keywords['medium']):
            grading_results['task_complexity_multiplier'] = 2.0
        else:
            grading_results['task_complexity_multiplier'] = 1.0
        
        # 3. Skill Adherence (bonus/penalty for correct tool usage)
        # In RAG training, this would check if appropriate knowledge retrieval was used
        if kb_results and len(kb_results) > 0:
            grading_results['skill_adherence'] = 100  # Used RAG correctly
        else:
            grading_results['skill_adherence'] = 50   # Partial credit
        
        # 4. Time Efficiency (bonus for faster completion)
        # This would be measured in real implementation
        response_length = len(response)
        if response_length > 100 and response_length < 1000:
            grading_results['time_efficiency'] = 100  # Good length
        elif response_length < 50:
            grading_results['time_efficiency'] = 50   # Too short
        else:
            grading_results['time_efficiency'] = 80   # Acceptable
        
        # 5. Calculate Overall Score using THE_ANSWER weights
        accuracy_weight = 0.4
        skill_weight = 0.3
        efficiency_weight = 0.3
        
        base_score = (
            grading_results['task_accuracy'] * accuracy_weight +
            grading_results['skill_adherence'] * skill_weight +
            grading_results['time_efficiency'] * efficiency_weight
        )
        
        # Apply complexity multiplier
        grading_results['overall_score'] = base_score * grading_results['task_complexity_multiplier']
        
        # Clamp to 0-100 range
        grading_results['overall_score'] = min(max(grading_results['overall_score'], 0), 100)
        
        print(f"   ✅ THE_ANSWER Grading Results:")
        print(f"      Task Accuracy: {grading_results['task_accuracy']:.1f}%")
        print(f"      Complexity Multiplier: ×{grading_results['task_complexity_multiplier']}")
        print(f"      Skill Adherence: {grading_results['skill_adherence']:.1f}%")
        print(f"      Time Efficiency: {grading_results['time_efficiency']:.1f}%")
        print(f"      Overall Score: {grading_results['overall_score']:.1f}%")
        
        return grading_results
    
    def _compress_context(self, kb_results: List[Dict], query: str) -> str:
        """Compress knowledge base context to most relevant parts"""
        if not kb_results:
            return ""
        
        # Simple compression: take first 500 chars of each result
        compressed_parts = []
        for result in kb_results[:3]:  # Limit to top 3 results
            document = result.get('document', '')
            if len(document) > 500:
                compressed_parts.append(document[:500] + "...")
            else:
                compressed_parts.append(document)
        
        return "\n\n".join(compressed_parts)


# Singleton instance
_rag_training_service = None

def get_rag_training_service() -> RAGTrainingService:
    """Get or create RAGTrainingService singleton"""
    global _rag_training_service
    if _rag_training_service is None:
        _rag_training_service = RAGTrainingService()
    return _rag_training_service
